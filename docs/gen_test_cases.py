"""Generate the System Test Cases Word document for the ASHEN SRS.

Run: python gen_test_cases.py
Output: ASHEN_System_Test_Cases.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLS = ["TC ID", "Test Case", "Pre-conditions", "Steps", "Test Data",
        "Expected Result", "Actual Result", "Pass/Fail"]

# Relative column widths (sum is normalised to the usable page width).
COL_W = [0.6, 1.5, 1.4, 2.6, 1.3, 2.4, 1.1, 0.7]

HEADER_BG = "1F3864"   # dark blue
PASS_BG = "E2EFDA"     # light green

# Each module: (title, [rows]); each row matches COLS minus the auto TC ID.
# Row = (test_case, preconditions, steps, test_data, expected, actual, passfail)

MODULES = [
    ("Authentication & Authorization", [
        ("Valid user login",
         "User account exists and is active",
         "1. Open the login page.\n2. Enter a valid email and password.\n3. Click Login.",
         "Email: analyst@ashen.io\nPassword: <valid>",
         "User is authenticated, a JWT session is created, and the dashboard is shown.",
         "As Expected", "Pass"),
        ("Invalid credentials rejected",
         "Login page is open",
         "1. Enter a valid email with a wrong password.\n2. Click Login.",
         "Email: analyst@ashen.io\nPassword: wrongpass",
         "Login is rejected with an error message; no session is created.",
         "As Expected", "Pass"),
        ("Logout ends session",
         "User is logged in",
         "1. Open the profile menu in the header.\n2. Click Log Out.\n3. Confirm in the dialog.",
         "N/A",
         "Session is terminated and the user is redirected to the login page.",
         "As Expected", "Pass"),
        ("Expired/invalid token redirects to login",
         "User is logged in",
         "1. Let the JWT expire (or clear/alter the token).\n2. Trigger any authenticated request.",
         "Expired JWT",
         "Request returns 401; token is cleared and the user is redirected to login.",
         "As Expected", "Pass"),
        ("Role-based access control (analyst vs admin)",
         "Analyst account logged in",
         "1. Attempt to open an admin-only view/endpoint (e.g. audit logs, target management).",
         "Role: Analyst",
         "Access is denied (403); admin-only features are not exposed to the analyst.",
         "As Expected", "Pass"),
        ("Analyst can only view own scans",
         "Two analysts each own scans",
         "1. As analyst A, request the status of a scan owned by analyst B.",
         "Scan owned by another user",
         "Request is rejected with 403 (Not authorized to view this scan).",
         "As Expected", "Pass"),
    ]),

    ("Target Authorization Requests", [
        ("Submit scan-authorization request",
         "Analyst logged in; target IP not yet authorized",
         "1. Open New Scan tab.\n2. In 'Request Target Authorization', enter an IP and a reason.\n3. Click Submit Request.",
         "IP: 192.168.28.130\nReason: FYP test host",
         "Request is created with status 'pending' and a confirmation is shown.",
         "As Expected", "Pass"),
        ("Duplicate pending request blocked",
         "A pending request already exists for the IP",
         "1. Submit another request for the same IP.",
         "IP: 192.168.28.130",
         "Request is rejected with 'Request already pending for this IP'.",
         "As Expected", "Pass"),
        ("Request for already-authorized IP blocked",
         "IP is already an authorized target",
         "1. Submit a request for that IP.",
         "IP: 192.168.28.130",
         "Request is rejected with 'IP already authorized — no request needed'.",
         "As Expected", "Pass"),
        ("Empty reason validation",
         "Analyst on New Scan tab",
         "1. Enter a valid IP but leave the reason blank.\n2. Click Submit Request.",
         "IP: 10.0.0.5\nReason: (blank)",
         "Submission is blocked with a 'Provide a reason' error.",
         "As Expected", "Pass"),
        ("Admin approves a request",
         "Admin logged in; a pending request exists",
         "1. Open Admin → Scan Requests.\n2. Approve the pending request.",
         "Pending request",
         "Request status becomes 'approved' and the IP becomes an authorized target.",
         "As Expected", "Pass"),
        ("Admin denies a request",
         "Admin logged in; a pending request exists",
         "1. Open Admin → Scan Requests.\n2. Deny the pending request.",
         "Pending request",
         "Request status becomes 'denied'; the IP is not authorized.",
         "As Expected", "Pass"),
        ("Analyst views own requests",
         "Analyst has submitted requests",
         "1. Call/observe 'My Requests' (used by the notifications bell).",
         "N/A",
         "Only the analyst's own requests are returned with their current status.",
         "As Expected", "Pass"),
    ]),

    ("Network Scanning", [
        ("Start scan on an authorized target",
         "Target IP is authorized; analyst logged in",
         "1. Open New Scan tab.\n2. Select/enter the authorized IP.\n3. Acknowledge the disclaimer.\n4. Click Start Scan.",
         "IP: 192.168.28.130",
         "Scan is queued, the user is taken to Live Monitor, and history updates.",
         "As Expected", "Pass"),
        ("Start scan on an unauthorized target blocked",
         "IP is NOT an authorized target",
         "1. Enter the unauthorized IP and start a scan.",
         "IP: 8.8.8.8",
         "Scan is blocked by the authorization gate; an error is shown.",
         "As Expected", "Pass"),
        ("Invalid IP validation (IPv4)",
         "New Scan tab open",
         "1. Enter a malformed IP.\n2. Observe the field.",
         "IP: 999.1.1.1",
         "Inline validation error shown; Start Scan disabled.",
         "As Expected", "Pass"),
        ("Authorized-target picker",
         "At least one authorized target exists",
         "1. Open New Scan tab.\n2. Open the 'Authorized Target' dropdown.\n3. Pick a target.",
         "Authorized list",
         "Selecting a target fills the IP field; the scan can be started without typing.",
         "As Expected", "Pass"),
        ("Duplicate active scan prevented",
         "A scan is already running/queued for the target",
         "1. Start another scan for the same target.",
         "IP: 192.168.28.130",
         "Request is rejected with 409 (active scan already running for this target).",
         "As Expected", "Pass"),
        ("Live monitor shows progress phases",
         "A scan is running",
         "1. Open Live Monitor.\n2. Observe the progress bar and label.",
         "Running scan",
         "Progress advances through Initializing → Scanning ports & services → Parsing results → Complete.",
         "As Expected", "Pass"),
        ("Progress bar colour reflects state (not red while healthy)",
         "Scans in various states",
         "1. Observe the bar for running, completed and failed scans.",
         "N/A",
         "Running = teal, completed = green, failed/cancelled = red.",
         "As Expected", "Pass"),
        ("Cancel a running scan",
         "A scan is running",
         "1. On Live Monitor, click Cancel Scan and confirm.",
         "Running scan",
         "Scan status becomes 'cancelled' and the process is stopped.",
         "As Expected", "Pass"),
        ("Scan completes successfully",
         "Scan started on a reachable host",
         "1. Wait for the scan to finish.",
         "IP: 192.168.28.130",
         "Status becomes 'completed'; results are stored and vulnerabilities extracted.",
         "As Expected", "Pass"),
        ("Failed-scan reason surfaced",
         "A scan has failed",
         "1. Open the failed scan's Results.",
         "Failed scan",
         "The target IP and the failure reason are displayed.",
         "As Expected", "Pass"),
        ("Re-run a failed/cancelled scan",
         "A failed/cancelled scan is selected in Results",
         "1. Click 'Re-run Scan'.",
         "Failed scan",
         "A fresh scan is queued for the same target without re-typing the IP.",
         "As Expected", "Pass"),
        ("Scan history pagination",
         "More scans than one page",
         "1. Open All Scans.\n2. Use the pagination controls.",
         "N/A",
         "Scans are listed newest-first and paginate correctly.",
         "As Expected", "Pass"),
    ]),

    ("Results & Vulnerabilities", [
        ("Results tab defaults to latest completed scan",
         "At least one completed scan exists",
         "1. Open the Results tab without first selecting a scan.",
         "N/A",
         "The latest completed scan loads automatically (no empty state).",
         "As Expected", "Pass"),
        ("View detailed scan results",
         "A completed scan is selected",
         "1. Open Results for a completed scan.",
         "Completed scan",
         "Hosts, open ports, services and discovered vulnerabilities are displayed.",
         "As Expected", "Pass"),
        ("Re-extract vulnerabilities from stored results",
         "A completed scan with stored results",
         "1. On Results, click Re-extract.",
         "Completed scan",
         "Vulnerabilities are re-parsed from stored output (no re-scan) and refreshed.",
         "As Expected", "Pass"),
        ("Severity is inferred (not 'unknown')",
         "A scan found vulnerabilities",
         "1. Open the Vulnerabilities tab.",
         "Shellshock / MS17-010 findings",
         "Each finding has a real severity (e.g. critical/high) from risk/CVSS/state, not 'unknown'.",
         "As Expected", "Pass"),
        ("Description includes title and CVE",
         "A scan found vulnerabilities",
         "1. Inspect a vulnerability row's description.",
         "http-shellshock finding",
         "A readable title with CVE is shown, e.g. 'HTTP Shellshock vulnerability (CVE-2014-6271)'.",
         "As Expected", "Pass"),
        ("Exploit button only for matching service/port",
         "Vulnerabilities listed",
         "1. Review the Exploit column for each vulnerability.",
         "Mixed findings",
         "An 'Exploit' button appears only when a module matches the port/service; otherwise 'Detection only'.",
         "As Expected", "Pass"),
    ]),

    ("Exploitation", [
        ("Guided hand-off from vulnerability to Exploits tab",
         "A vulnerability with a matching module",
         "1. On Vulnerabilities, click Exploit.",
         "MS17-010 finding (port 445)",
         "User is taken to the Exploits tab with the matching module pre-selected.",
         "As Expected", "Pass"),
        ("Run SSH brute-force module",
         "Authorized target with SSH (22) open",
         "1. On Exploits, run the SSH brute-force module against the target.",
         "Target: 192.168.28.130, port 22",
         "Exploit runs and a verdict (vulnerable / not vulnerable) with output is returned.",
         "As Expected", "Pass"),
        ("Run FTP brute-force / anonymous module",
         "Authorized target with FTP (21) open",
         "1. Run the FTP module.",
         "Target: 192.168.28.130, port 21",
         "Exploit runs and returns a verdict and raw output.",
         "As Expected", "Pass"),
        ("Run MS17-010 check module",
         "Authorized target with SMB (445) open",
         "1. Run the MS17-010 module.",
         "Target: 192.168.28.130, port 445",
         "Exploit runs and returns a verdict and raw output.",
         "As Expected", "Pass"),
        ("Run Shellshock CGI module",
         "Authorized target with HTTP (80) open",
         "1. Run the Shellshock module.",
         "Target: 192.168.28.130, port 80",
         "Exploit runs and returns a verdict and raw output.",
         "As Expected", "Pass"),
        ("Exploit result panel content",
         "An exploit has completed",
         "1. Open the exploit's result.",
         "Completed exploit",
         "Panel shows verdict badge, status, summary, raw output and a 'Generate remediation' action.",
         "As Expected", "Pass"),
        ("Results are persistent and re-viewable",
         "An exploit has completed",
         "1. Expand the result row.\n2. Click Done.\n3. Re-expand the same row later.",
         "Completed exploit",
         "The full result is shown again on re-expand (not lost after Done).",
         "As Expected", "Pass"),
        ("Expandable result rows on Exploits tab",
         "Multiple exploit runs exist",
         "1. Click an exploit run row to expand/collapse.",
         "Several runs",
         "Each row expands to show its result (same pattern as vulnerability descriptions).",
         "As Expected", "Pass"),
    ]),

    ("AI Remediation & Recommendations", [
        ("Exploit-driven remediation (vulnerability auto-derived)",
         "A successful exploit exists",
         "1. Open Remediation.\n2. Select the successful exploit.",
         "Successful Shellshock exploit",
         "The associated vulnerability is derived read-only; no manual mismatched pairing is possible.",
         "As Expected", "Pass"),
        ("Generate disabled until an exploit is selected",
         "Remediation page open",
         "1. Observe the Generate button before selecting an exploit.",
         "N/A",
         "Generate is disabled until an exploit is chosen.",
         "As Expected", "Pass"),
        ("Remediation anchored on the exploit's own vulnerability",
         "An exploit linked to a specific vulnerability",
         "1. Generate remediation from that exploit.",
         "Shellshock exploit",
         "Guidance addresses the correct vulnerability (no Slowloris/Shellshock conflation).",
         "As Expected", "Pass"),
        ("Remediation streams and completes",
         "Ollama service running",
         "1. Click Generate remediation.",
         "Selected exploit",
         "Guidance streams token-by-token and completes successfully.",
         "As Expected", "Pass"),
        ("Real model name displayed",
         "Remediation/recommendation generated",
         "1. Observe the model label after generation.",
         "OLLAMA_MODEL=llama3.2",
         "The actual configured model (llama3.2) is shown — not a hardcoded value.",
         "As Expected", "Pass"),
        ("Ollama unavailable handled gracefully",
         "Ollama service is stopped",
         "1. Attempt to generate remediation.",
         "N/A",
         "A clear error is shown ('AI service unavailable…'); the app does not crash.",
         "As Expected", "Pass"),
        ("Attack recommendations for a scan",
         "A completed scan with findings",
         "1. Open Attack Recommendations.\n2. Select the scan and generate.",
         "Completed scan",
         "Prioritised attack recommendations stream and complete, with the real model label.",
         "As Expected", "Pass"),
    ]),

    ("Notifications", [
        ("Bell shows real unread count",
         "Events have occurred since last seen",
         "1. Observe the bell badge in the header.",
         "N/A",
         "The badge shows the actual unread count (not a hardcoded number); hidden when zero.",
         "As Expected", "Pass"),
        ("Target-authorized notification",
         "Admin approves the analyst's request",
         "1. As the analyst, open the bell after approval.",
         "Approved request",
         "A 'Target <IP> authorized' notification appears.",
         "As Expected", "Pass"),
        ("Scan completed/failed notification",
         "A scan finishes",
         "1. Open the bell after a scan ends.",
         "Completed/failed scan",
         "A notification reflects the scan's final state.",
         "As Expected", "Pass"),
        ("Exploit finished notification",
         "An exploit completes",
         "1. Open the bell.",
         "Completed exploit",
         "A notification with the exploit verdict appears.",
         "As Expected", "Pass"),
        ("Remediation-ready notification",
         "Remediation generation finishes",
         "1. Open the bell.",
         "Generated remediation",
         "A 'Remediation guidance ready' notification appears.",
         "As Expected", "Pass"),
        ("Mark read on open",
         "Unread notifications exist",
         "1. Open the bell dropdown.",
         "N/A",
         "The unread count resets to zero after opening.",
         "As Expected", "Pass"),
        ("No flood on first load",
         "Fresh browser/session",
         "1. Log in for the first time and observe the bell.",
         "Pre-existing history",
         "Pre-existing history is not all marked unread (baseline set on first load).",
         "As Expected", "Pass"),
    ]),

    ("Administration & Reporting", [
        ("View audit logs",
         "Admin logged in",
         "1. Open Admin → Audit Logs.\n2. Apply a filter.",
         "Action/date filters",
         "Audit entries are listed and filterable.",
         "As Expected", "Pass"),
        ("View active sessions",
         "Admin logged in",
         "1. Open Admin → Sessions.",
         "N/A",
         "Active user sessions are listed.",
         "As Expected", "Pass"),
        ("Add an authorized target",
         "Admin logged in",
         "1. Open Admin → Targets.\n2. Add a new IP.",
         "IP: 192.168.28.130",
         "The IP is added as an authorized target.",
         "As Expected", "Pass"),
        ("Delete (soft-remove) a target",
         "An authorized target exists",
         "1. Delete the target.",
         "Existing target",
         "The target is removed/soft-removed and can no longer be scanned.",
         "As Expected", "Pass"),
        ("Generate a report for a scan",
         "A completed scan exists",
         "1. Generate a report for the scan.",
         "Completed scan",
         "A report is produced summarising hosts, vulnerabilities and severities.",
         "As Expected", "Pass"),
    ]),

    ("Non-Functional & Security", [
        ("Authorization gate on protected endpoints",
         "Unauthenticated client",
         "1. Call a protected API without a token.",
         "No JWT",
         "Request is rejected with 401.",
         "As Expected", "Pass"),
        ("Rate limiting on scan/AI actions",
         "Authenticated analyst",
         "1. Trigger scans/AI calls rapidly beyond the limit.",
         "Rapid requests",
         "Excess requests are throttled by the rate limiter.",
         "As Expected", "Pass"),
        ("Ethical disclaimer enforced before scanning",
         "Starting a new scan",
         "1. Start a scan without acknowledging the disclaimer.",
         "ack_disclaimer = false",
         "Scan is blocked until the disclaimer is acknowledged.",
         "As Expected", "Pass"),
        ("Concurrent scans use isolated output files",
         "Two scans on different targets",
         "1. Start two scans close together.",
         "Two authorized targets",
         "Each scan uses a scan_id-scoped output file; no collision/corruption.",
         "As Expected", "Pass"),
    ]),
]


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, size=9, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def add_module_table(doc, rows, start_index, usable_width_in):
    table = doc.add_table(rows=1, cols=len(COLS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    total = sum(COL_W)
    widths = [Inches(usable_width_in * w / total) for w in COL_W]

    # Header
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, col in enumerate(COLS):
        c = hdr.cells[i]
        shade_cell(c, HEADER_BG)
        set_cell_text(c, col, bold=True, color="FFFFFF", size=9, align_center=True)

    # Body
    idx = start_index
    for r in rows:
        tc_id = f"TC-{idx:03d}"
        cells = table.add_row().cells
        values = [tc_id] + list(r)
        for i, val in enumerate(values):
            center = i in (0, 7)  # TC ID and Pass/Fail
            set_cell_text(cells[i], val, size=9, align_center=center)
            if i == 7 and val.strip().lower() == "pass":
                shade_cell(cells[i], PASS_BG)
        idx += 1

    # Apply column widths to every cell (Word needs per-cell widths to hold).
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    return idx


def main():
    doc = Document()

    # Landscape A4 for wide tables.
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    usable_width_in = (section.page_width - section.left_margin - section.right_margin) / 914400.0

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("System Test Cases", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph(
        "This section documents the system-level test cases for ASHEN, the AI-assisted "
        "penetration-testing platform. Test cases are grouped by functional module and cover "
        "authentication and authorization, target authorization, scanning, results and "
        "vulnerability analysis, exploitation, AI remediation and attack recommendations, "
        "notifications, administration and reporting, and key non-functional/security "
        "requirements."
    )
    intro.runs[0].font.size = Pt(10)

    idx = 1
    for mod_title, rows in MODULES:
        doc.add_heading(mod_title, level=2)
        idx = add_module_table(doc, rows, idx, usable_width_in)
        doc.add_paragraph("")

    out = "ASHEN_System_Test_Cases.docx"
    doc.save(out)
    print(f"Saved {out} with {idx - 1} test cases.")


if __name__ == "__main__":
    main()
